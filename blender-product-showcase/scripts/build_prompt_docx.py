#!/usr/bin/env python3
"""Build a user-editable DOCX from references/prompt-template.md."""

from __future__ import annotations

import argparse
from pathlib import Path


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = SKILL_ROOT / "references" / "prompt-template.md"
DEFAULT_OUTPUT = SKILL_ROOT / "assets" / "Blender产品展示通用提示词模板.docx"


def build(source: Path, output: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit("python-docx is required: pip install python-docx") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    in_code = False
    code_lines: list[str] = []
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines.clear()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line and line[0].isdigit() and ". " in line[:4]:
            document.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line == "---":
            document.add_page_break()
        elif line:
            document.add_paragraph(line)
        else:
            document.add_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
